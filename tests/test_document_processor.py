import os

import pytest
from bs4 import BeautifulSoup

from utils.langchain.document_processor import DocumentProcessor


def test_process_text():
    dp = DocumentProcessor()
    content = "This is a simple test."
    documents = dp.process_text(content)
    assert len(documents) > 0
    assert content in documents[0].page_content


def test_process_html(tmp_path):
    dp = DocumentProcessor()
    html_content = "<html><body><h1>Title</h1><p>Paragraph.</p></body></html>"
    html_file = tmp_path / "test.html"
    html_file.write_text(html_content, encoding="utf-8")
    documents = dp.process_html(html_file)
    # Use BeautifulSoup to verify extraction
    soup = BeautifulSoup(html_content, "html.parser")
    text = soup.get_text(separator="\n")
    assert any(text.strip() in doc.page_content for doc in documents)


@pytest.mark.skipif(os.name == "nt" and "USERPROFILE" in os.environ, reason="Requires python-docx file creation")
def test_process_docx(tmp_path):
    from docx import Document as DocxDocument

    dp = DocumentProcessor()
    docx_file = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("Paragraph one.")
    doc.add_paragraph("Paragraph two.")
    doc.save(docx_file)
    documents = dp.process_docx(docx_file)
    combined = " ".join(d.page_content for d in documents)
    assert "Paragraph one." in combined
    assert "Paragraph two." in combined
